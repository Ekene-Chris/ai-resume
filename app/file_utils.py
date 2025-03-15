# app/file_utils.py
import io
import os
import logging
from typing import Optional, Union
import aiohttp
import asyncio
from urllib.parse import urlparse

# Try to import PDF and DOCX handling libraries
try:
    import PyPDF2
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False
    logging.error("PyPDF2 library not installed. PDF support is limited.")

try:
    import docx
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False
    logging.error("python-docx library not installed. DOCX support is limited.")

logger = logging.getLogger(__name__)

async def extract_text_from_pdf_bytes(pdf_bytes: bytes) -> str:
    """
    Extract text from PDF bytes using PyPDF2
    
    Args:
        pdf_bytes: PDF content as bytes
        
    Returns:
        Extracted text
    """
    if not PDF_SUPPORT:
        logger.error("PDF extraction failed: PyPDF2 library not installed")
        # Try a simplistic fallback method
        return fallback_binary_to_text(pdf_bytes, "PDF")
    
    try:
        text = ""
        pdf_file = io.BytesIO(pdf_bytes)
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        
        # Log PDF details
        logger.info(f"PDF contains {len(pdf_reader.pages)} pages")
        
        for page_num in range(len(pdf_reader.pages)):
            page = pdf_reader.pages[page_num]
            try:
                page_text = page.extract_text()
                text += page_text + "\n\n"
                
                # Log empty pages
                if not page_text.strip():
                    logger.warning(f"Page {page_num + 1} appears to be empty or contains no extractable text")
            except Exception as page_error:
                logger.warning(f"Error extracting text from page {page_num + 1}: {str(page_error)}")
        
        # Log if extracted text is too short
        if len(text.strip()) < 100:
            logger.warning(f"Extracted PDF text is suspiciously short ({len(text.strip())} chars)")
            
        return text
    except Exception as e:
        logger.error(f"Error extracting text from PDF: {str(e)}", exc_info=True)
        # Try fallback method
        return fallback_binary_to_text(pdf_bytes, "PDF")

async def extract_text_from_docx_bytes(docx_bytes: bytes) -> str:
    """
    Extract text from DOCX bytes using python-docx
    
    Args:
        docx_bytes: DOCX content as bytes
        
    Returns:
        Extracted text
    """
    if not DOCX_SUPPORT:
        logger.error("DOCX extraction failed: python-docx library not installed")
        # Try a simplistic fallback method
        return fallback_binary_to_text(docx_bytes, "DOCX")
    
    try:
        text = ""
        docx_file = io.BytesIO(docx_bytes)
        doc = docx.Document(docx_file)
        
        # Log document details
        logger.info(f"DOCX contains {len(doc.paragraphs)} paragraphs and {len(doc.tables)} tables")
        
        for para in doc.paragraphs:
            text += para.text + "\n"
        
        # Also extract tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " | "
                text += "\n"
        
        # Log if extracted text is too short
        if len(text.strip()) < 100:
            logger.warning(f"Extracted DOCX text is suspiciously short ({len(text.strip())} chars)")
            
        return text
    except Exception as e:
        logger.error(f"Error extracting text from DOCX: {str(e)}", exc_info=True)
        # Try fallback method
        return fallback_binary_to_text(docx_bytes, "DOCX")

def fallback_binary_to_text(binary_data: bytes, file_type: str) -> str:
    """
    Attempt to extract text from binary data when specialized libraries fail
    
    Args:
        binary_data: The binary file data
        file_type: The type of file ("PDF", "DOCX", etc.)
        
    Returns:
        Extracted text or error message
    """
    logger.info(f"Using fallback text extraction for {file_type}")
    
    try:
        # Try to decode as UTF-8 with error handling
        text = binary_data.decode('utf-8', errors='replace')
        
        # Clean up the text - remove unprintable characters
        cleaned_text = ''.join(char if ord(char) >= 32 or char == '\n' else ' ' for char in text)
        
        # Make sure we actually got some usable text
        if len(cleaned_text.strip()) > 50:
            return cleaned_text
        
        logger.warning(f"Fallback text extraction produced minimal text: {len(cleaned_text.strip())} chars")
        return f"[Document content could not be fully extracted. Please ensure your {file_type} file contains selectable text.]"
    except Exception as e:
        logger.error(f"Fallback text extraction failed: {str(e)}")
        return f"[Document content could not be extracted from {file_type}. Error: {str(e)}]"

async def download_file(url: str) -> Optional[bytes]:
    """
    Download a file from a URL with improved error handling
    
    Args:
        url: The URL to download from
        
    Returns:
        The file content as bytes, or None if download fails
    """
    # Validate URL
    try:
        parsed_url = urlparse(url)
        if not all([parsed_url.scheme, parsed_url.netloc]):
            logger.error(f"Invalid URL format: {url}")
            return None
    except Exception as url_error:
        logger.error(f"URL parsing error: {str(url_error)}")
        return None
    
    # Set up timeout and retries
    timeout = aiohttp.ClientTimeout(total=60)  # 60 seconds total timeout
    max_retries = 3
    
    # Create custom headers - some storage services need user agent
    headers = {
        "User-Agent": "ResumeScannerAgent/1.0 (Teleios/EkeneChris Resume Analysis Service)"
    }
    
    for attempt in range(max_retries):
        try:
            logger.info(f"Downloading file from URL (attempt {attempt+1}/{max_retries}): {url}")
            
            async with aiohttp.ClientSession(timeout=timeout) as session:
                async with session.get(url, headers=headers) as response:
                    if response.status == 200:
                        content = await response.read()
                        content_size = len(content)
                        
                        logger.info(f"Successfully downloaded file: {content_size} bytes")
                        
                        if content_size < 100:
                            logger.warning(f"Downloaded file is suspiciously small: {content_size} bytes")
                        
                        return content
                    else:
                        error_text = await response.text()
                        error_truncated = error_text[:200] + ("..." if len(error_text) > 200 else "")
                        
                        logger.error(f"Failed to download file: HTTP {response.status}")
                        logger.error(f"Error response: {error_truncated}")
                        
                        # HTTP 403 or 404 - no point retrying
                        if response.status in [403, 404]:
                            logger.error(f"Resource not accessible (HTTP {response.status}), not retrying")
                            raise Exception(f"File not accessible: HTTP {response.status}")
                        
                        # Wait before retrying
                        if attempt < max_retries - 1:
                            await asyncio.sleep(2 ** attempt)  # Exponential backoff
        
        except asyncio.TimeoutError:
            logger.error(f"Timeout downloading file (attempt {attempt+1}/{max_retries})")
            if attempt < max_retries - 1:
                await asyncio.sleep(2 ** attempt)  # Exponential backoff
        
        except Exception as e:
            logger.error(f"Error downloading file (attempt {attempt+1}/{max_retries}): {str(e)}", exc_info=True)
            if attempt < max_retries - 1:
                await asyncio.sleep(2 ** attempt)  # Exponential backoff
    
    logger.error(f"All download attempts failed for URL: {url}")
    return None

async def get_file_details(file_bytes: bytes, filename: str) -> dict:
    """
    Get details about a file for logging and debugging
    
    Args:
        file_bytes: The file content
        filename: The name of the file
        
    Returns:
        Dict with file details
    """
    file_ext = os.path.splitext(filename)[1].lower() if '.' in filename else ''
    
    details = {
        "filename": filename,
        "file_extension": file_ext,
        "file_size": len(file_bytes),
        "content_sample": file_bytes[:50].hex() if file_bytes else None,
    }
    
    # Try to get more details based on file type
    if file_ext == '.pdf' and PDF_SUPPORT:
        try:
            pdf_file = io.BytesIO(file_bytes)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            details["page_count"] = len(pdf_reader.pages)
            details["metadata"] = pdf_reader.metadata
        except Exception as e:
            details["pdf_error"] = str(e)
    
    elif file_ext in ['.doc', '.docx'] and DOCX_SUPPORT:
        try:
            docx_file = io.BytesIO(file_bytes)
            doc = docx.Document(docx_file)
            details["paragraph_count"] = len(doc.paragraphs)
            details["table_count"] = len(doc.tables)
        except Exception as e:
            details["docx_error"] = str(e)
    
    return details